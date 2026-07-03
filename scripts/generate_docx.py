#!/usr/bin/env python3
"""
生成专业的解决方案 DOCX 文档

将 Markdown 格式的解决方案内容转换为排版精美的 DOCX 文档，
包含封面、目录、页眉页脚等专业排版元素，并自动去除 Markdown 格式符号。
"""

import os
import sys
import argparse
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def clean_markdown_format(text):
    """
    去除 Markdown 格式符号
    
    Args:
        text: 包含 Markdown 格式的文本
    
    Returns:
        清理后的纯文本
    """
    if not text:
        return text
    
    # 去除加粗 **text**
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # 去除斜体 *text*
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # 去除删除线 ~~text~~
    text = re.sub(r'~~(.*?)~~', r'\1', text)
    
    # 去除行内代码 `text`
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # 去除链接 [text](url)，保留 text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    
    # 去除无序列表标记（行首的 - 或 * 或 +）
    text = re.sub(r'^[\s]*[-*+][\s]+', '', text, flags=re.MULTILINE)
    
    # 去除有序列表标记（行首的 1. 或 2. 等）
    text = re.sub(r'^[\s]*\d+\.[\s]+', '', text, flags=re.MULTILINE)
    
    # 去除标题标记（行首的 #）
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    
    # 去除分隔线（--- 或 *** 等）
    text = re.sub(r'^[\s]*[-*_]{3,}[\s]*$', '', text, flags=re.MULTILINE)
    
    # 去除引用标记（行首的 >）
    text = re.sub(r'^[\s]*>[\s]*', '', text, flags=re.MULTILINE)
    
    # 去除多余的空行（连续3个及以上换行符改为2个）
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text


def set_chinese_font(run, font_name="宋体", size=11):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)


def set_body_paragraph_format(paragraph):
    """统一正文段落样式：首行缩进2字符、无段前后空白。"""
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(22)  # 约等于 11pt 字体下首行缩进2字符
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def setup_page_layout(doc):
    """设置页面版式，采用常见商务文档参数。"""
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def setup_styles(doc):
    """设置全局样式，确保输出稳定且统一。"""
    normal_style = doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    heading1 = doc.styles['Heading 1']
    heading1.font.name = '黑体'
    heading1.font.size = Pt(18)
    heading1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.line_spacing = 1.3

    heading2 = doc.styles['Heading 2']
    heading2.font.name = '黑体'
    heading2.font.size = Pt(16)
    heading2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(8)
    heading2.paragraph_format.line_spacing = 1.3

    heading3 = doc.styles['Heading 3']
    heading3.font.name = '黑体'
    heading3.font.size = Pt(14)
    heading3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(6)
    heading3.paragraph_format.line_spacing = 1.25


def create_solution_docx(content, output_path, project_name, customer_name, input_base_dir):
    """
    创建解决方案 DOCX 文档
    
    Args:
        content: 文档内容（Markdown格式）
        output_path: 输出文件路径
        project_name: 项目名称
        customer_name: 客户名称
        input_base_dir: 输入Markdown所在目录（用于解析图片相对路径）
    """
    doc = Document()

    # 设置文档版式与样式
    setup_page_layout(doc)
    setup_styles(doc)
    
    # 1. 添加封面
    add_cover_page(doc, project_name, customer_name)
    
    # 2. 添加分页符
    doc.add_page_break()
    
    # 3. 添加目录占位符
    add_toc_placeholder(doc)
    
    # 4. 添加分页符
    doc.add_page_break()
    
    # 5. 解析并添加正文内容
    add_body_content(doc, content, input_base_dir)
    
    # 6. 保存文档
    doc.save(output_path)
    print(f"文档已生成: {output_path}")
    return output_path


def add_cover_page(doc, project_name, customer_name):
    """添加封面页"""
    # 封面标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加空行
    for _ in range(8):
        doc.add_paragraph()
    
    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{customer_name}")
    set_chinese_font(run, "黑体", 22)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{project_name}")
    set_chinese_font(run, "黑体", 28)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("解决方案")
    set_chinese_font(run, "黑体", 32)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 添加空行
    for _ in range(6):
        doc.add_paragraph()
    
    # 底部信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}")
    set_chinese_font(run, "宋体", 14)


def add_toc_placeholder(doc):
    """添加目录占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目  录")
    set_chinese_font(run, "黑体", 18)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ 目录将在打开文档时自动生成 ]")
    set_chinese_font(run, "宋体", 11)
    run.font.color.rgb = RGBColor(128, 128, 128)


def parse_markdown_table(lines, start_index):
    """
    解析 Markdown 表格
    
    Args:
        lines: 所有行
        start_index: 表格开始的行索引
    
    Returns:
        (table_data, end_index): 表格数据和结束行索引
    """
    table_data = []
    i = start_index
    
    # 读取表头
    if i < len(lines) and lines[i].startswith('|'):
        header_row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
        # 清理表头中的 Markdown 格式
        header_row = [clean_markdown_format(cell) for cell in header_row]
        table_data.append(header_row)
        i += 1
    
    # 跳过分隔行（|---|---|）
    if i < len(lines) and lines[i].startswith('|') and '---' in lines[i]:
        i += 1
    
    # 读取数据行
    while i < len(lines) and lines[i].startswith('|'):
        data_row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
        # 清理单元格中的 Markdown 格式
        data_row = [clean_markdown_format(cell) for cell in data_row]
        table_data.append(data_row)
        i += 1
    
    return table_data, i


def parse_markdown_image(line):
    """解析 Markdown 图片语法：![alt](path)"""
    match = re.match(r'^\s*!\[(.*?)\]\((.*?)\)\s*$', line)
    if not match:
        return None, None
    alt = match.group(1).strip()
    path = match.group(2).strip()
    return alt, path


def resolve_image_path(image_path, base_dir):
    """解析图片路径，支持相对路径。"""
    p = Path(image_path)
    if p.is_absolute():
        return p
    return Path(base_dir) / p


def add_image_to_doc(doc, image_abs_path, caption, figure_index):
    """插入图片与图题。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(image_abs_path), width=Cm(15.5))

    caption_text = caption if caption else f"逻辑图{figure_index}"
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Pt(0)
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(10)
    cp.paragraph_format.line_spacing = 1.2
    c_run = cp.add_run(f"图{figure_index} {caption_text}")
    set_chinese_font(c_run, "宋体", 10)
    c_run.font.color.rgb = RGBColor(89, 89, 89)


def add_table_to_doc(doc, table_data):
    """将表格添加到文档"""
    if not table_data or len(table_data) < 2:
        return
    
    rows = len(table_data)
    cols = len(table_data[0])
    
    # 创建表格
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.autofit = True
    
    # 填充表格内容
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(cell_text)
            set_chinese_font(run, "宋体", 10)
            
            # 表头加粗
            if i == 0:
                run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_body_content(doc, content, input_base_dir):
    """添加正文内容"""
    lines = content.split('\n')
    i = 0
    figure_index = 1
    
    while i < len(lines):
        line = lines[i].rstrip()

        # 检查是否是图片
        alt_text, image_path = parse_markdown_image(line)
        if image_path:
            image_abs_path = resolve_image_path(image_path, input_base_dir)
            if image_abs_path.exists():
                add_image_to_doc(doc, image_abs_path, alt_text, figure_index)
                figure_index += 1
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(0)
                run = p.add_run(f"[图片未找到] {image_path}")
                set_chinese_font(run, "宋体", 10)
                run.font.color.rgb = RGBColor(160, 0, 0)
            i += 1
            continue
        
        # 检查是否是表格开始
        if line.startswith('|'):
            table_data, i = parse_markdown_table(lines, i)
            add_table_to_doc(doc, table_data)
            continue
        
        # 跳过空行，避免在 Word 中产生“段落间空一行”
        if not line:
            i += 1
            continue
        
        # 清理当前行的 Markdown 格式
        cleaned_line = clean_markdown_format(line)
        
        # 处理标题（保留标题层级，但去除 # 符号）
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.style = 'Heading 1'
            run = p.add_run(cleaned_line)
            set_chinese_font(run, "黑体", 18)
        
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.style = 'Heading 2'
            run = p.add_run(cleaned_line)
            set_chinese_font(run, "黑体", 16)
        
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.style = 'Heading 3'
            run = p.add_run(cleaned_line)
            set_chinese_font(run, "黑体", 14)
        
        else:
            # 普通段落
            p = doc.add_paragraph()
            run = p.add_run(cleaned_line)
            set_chinese_font(run, "宋体", 11)
            set_body_paragraph_format(p)
        
        i += 1


def main():
    parser = argparse.ArgumentParser(description='生成解决方案 DOCX 文档')
    parser.add_argument('--input', '-i', required=True, help='输入 Markdown 文件路径')
    parser.add_argument('--output', '-o', help='输出 DOCX 文件路径（可选）')
    parser.add_argument('--project', '-p', required=True, help='项目名称')
    parser.add_argument('--customer', '-c', required=True, help='客户名称')
    
    args = parser.parse_args()
    
    # 读取输入文件
    if not os.path.exists(args.input):
        print(f"错误：输入文件不存在: {args.input}")
        return 1
    
    try:
        with open(args.input, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        print(f"读取输入文件时出错: {str(e)}")
        return 1
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        # 生成默认文件名：YYYYMMDD_客户简称_项目名_方案.docx
        date_str = datetime.now().strftime('%Y%m%d')
        customer_short = args.customer[:4] if len(args.customer) > 4 else args.customer
        project_short = args.project[:6] if len(args.project) > 6 else args.project
        # 清理文件名中的非法字符
        customer_short = re.sub(r'[<>:"/\\|?*]', '', customer_short)
        project_short = re.sub(r'[<>:"/\\|?*]', '', project_short)
        output_path = f"{date_str}_{customer_short}_{project_short}_方案.docx"
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
        except Exception as e:
            print(f"创建输出目录时出错: {str(e)}")
            return 1
    
    # 生成文档
    try:
        input_base_dir = str(Path(args.input).resolve().parent)
        create_solution_docx(content, output_path, args.project, args.customer, input_base_dir)
        return 0
    except Exception as e:
        print(f"生成文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
